from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import hashlib
import re

from docx import Document


ARTICLE_RE = re.compile(r"^(第[一二三四五六七八九十百千万零〇两]+条)\s*(.*)$")
BOOK_RE = re.compile(r"^第[一二三四五六七八九十百千万零〇两]+编\s+")
CHAPTER_RE = re.compile(r"^第[一二三四五六七八九十百千万零〇两]+章\s+")
SECTION_RE = re.compile(r"^第[一二三四五六七八九十百千万零〇两]+节\s+")
PASS_DATE_RE = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日")


@dataclass(frozen=True)
class LawArticle:
    article_no: str
    article_order: int
    content: str
    book_title: str | None
    chapter_title: str | None
    section_title: str | None
    full_path: str
    content_hash: str


@dataclass(frozen=True)
class ParsedLawDocument:
    title: str
    law_type: str
    publish_date: str | None
    source_file_name: str
    articles: list[LawArticle]

    def article_by_no(self, article_no: str) -> LawArticle:
        for article in self.articles:
            if article.article_no == article_no:
                return article
        raise KeyError(article_no)


@dataclass
class _ArticleDraft:
    article_no: str
    article_order: int
    parts: list[str]
    book_title: str | None
    chapter_title: str | None
    section_title: str | None

    def append(self, text: str) -> None:
        if text:
            self.parts.append(text)

    def build(self, document_title: str) -> LawArticle:
        content = "\n".join(self.parts).strip()
        path_parts = [
            document_title,
            self.book_title,
            self.chapter_title,
            self.section_title,
            self.article_no,
        ]
        full_path = " > ".join(part for part in path_parts if part)
        return LawArticle(
            article_no=self.article_no,
            article_order=self.article_order,
            content=content,
            book_title=self.book_title,
            chapter_title=self.chapter_title,
            section_title=self.section_title,
            full_path=full_path,
            content_hash=hashlib.sha256(content.encode("utf-8")).hexdigest(),
        )


def parse_law_docx(path: str | Path, law_type: str = "法律") -> ParsedLawDocument:
    source_path = Path(path)
    paragraphs = _read_paragraphs(source_path)
    if len(paragraphs) < 3:
        raise ValueError(f"document has too few paragraphs: {source_path}")

    title = paragraphs[0]
    publish_date = _extract_publish_date(paragraphs[1])
    body_start = _find_body_start(paragraphs)

    book_title: str | None = None
    chapter_title: str | None = None
    section_title: str | None = None
    drafts: list[_ArticleDraft] = []
    current: _ArticleDraft | None = None

    for text in paragraphs[body_start:]:
        if BOOK_RE.match(text):
            book_title = _normalize_heading(text)
            chapter_title = None
            section_title = None
            continue
        if CHAPTER_RE.match(text):
            chapter_title = _normalize_heading(text)
            section_title = None
            continue
        if SECTION_RE.match(text):
            section_title = _normalize_heading(text)
            continue

        article_match = ARTICLE_RE.match(text)
        if article_match:
            article_no, article_content = article_match.groups()
            current = _ArticleDraft(
                article_no=article_no,
                article_order=len(drafts) + 1,
                parts=[article_content.strip()],
                book_title=book_title,
                chapter_title=chapter_title,
                section_title=section_title,
            )
            drafts.append(current)
            continue

        if current is not None:
            current.append(text)

    articles = [draft.build(title) for draft in drafts]
    return ParsedLawDocument(
        title=title,
        law_type=law_type,
        publish_date=publish_date,
        source_file_name=source_path.name,
        articles=articles,
    )


def _read_paragraphs(path: Path) -> list[str]:
    document = Document(path)
    return [_normalize_text(p.text) for p in document.paragraphs if _normalize_text(p.text)]


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def _normalize_heading(text: str) -> str:
    parts = text.split(" ", 1)
    if len(parts) == 1:
        return text
    prefix, title = parts
    return f"{prefix} {title.replace(' ', '')}"


def _extract_publish_date(text: str) -> str | None:
    match = PASS_DATE_RE.search(text)
    if not match:
        return None
    year, month, day = match.groups()
    return f"{int(year):04d}-{int(month):02d}-{int(day):02d}"


def _find_body_start(paragraphs: list[str]) -> int:
    first_article_index = next(
        (index for index, text in enumerate(paragraphs) if ARTICLE_RE.match(text)),
        None,
    )
    if first_article_index is None:
        raise ValueError("no article paragraphs found")

    start = first_article_index
    while start > 0 and _is_hierarchy_heading(paragraphs[start - 1]):
        start -= 1
    return start


def _is_hierarchy_heading(text: str) -> bool:
    return bool(BOOK_RE.match(text) or CHAPTER_RE.match(text) or SECTION_RE.match(text))
